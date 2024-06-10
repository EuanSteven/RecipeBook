# Copyright 2023 Euan Steven
# 
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# 
#     http://www.apache.org/licenses/LICENSE-2.0
# 
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

# -*- encoding: utf-8 -*-
# ========== export.py ==========
# Author: Euan Steven
# Date Created: 16/12/2023
# Date Modified: 17/12/2023
# Version: 1.7
# License: Apache 2.0
# Description: Export .txt File to .docx File
# ================================

# ========== Module Imports ==========

print("[OK] EXPORT - Importing Modules")

# Built-in Modules
import os

# Third-party Modules
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx2pdf import convert

print("[OK] EXPORT - Imported Modules")

# ========== Create Recipe Document ==========

def create_recipe_doc(recipe_name, ingredients, method):
    doc = Document()
    
    # Corrected Custom Styles
    title_style = doc.styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
    title_style.font.name = 'Times New Roman'
    title_style.font.size = Pt(16)
    title_style.font.bold = True

    heading_style = doc.styles.add_style('HeadingStyle', WD_STYLE_TYPE.PARAGRAPH)
    heading_style.font.name = 'Times New Roman'
    heading_style.font.size = Pt(14)
    heading_style.font.bold = True

    body_style = doc.styles.add_style('BodyStyle', WD_STYLE_TYPE.PARAGRAPH)
    body_style.font.name = 'Times New Roman'
    body_style.font.size = Pt(12)

    # Add Recipe Name as heading
    doc.add_heading(recipe_name, level=1).style = title_style
    
    # Add Ingredients
    doc.add_heading('Ingredients', level=2).style = heading_style
    for ingredient in ingredients:
        doc.add_paragraph(ingredient, style='BodyStyle')
    
    # Add Method
    doc.add_heading('Method', level=2).style = heading_style
    for step in method:
        doc.add_paragraph(step, style='BodyStyle')

    return doc

# ========== Read Recipe from File ==========

def read_recipe_from_file(file_path):
    try:
        with open(file_path, 'r') as file:
            lines = file.readlines()

        # Find indices for Recipe Name, Ingredients, and Method
        recipe_name_start = next((i for i, line in enumerate(lines) if line.startswith('Recipe: ')), None)
        ingredients_start = next((i for i, line in enumerate(lines) if 'Ingredients:' in line), None)
        method_start = next((i for i, line in enumerate(lines) if 'Method:' in line), None)

        # Check if the file has the expected content
        if None in (recipe_name_start, ingredients_start, method_start):
            raise ValueError("Invalid file format")

        # Extract Recipe Name
        recipe_name = lines[recipe_name_start].strip().split(": ")[1]

        # Extract Ingredients
        ingredients_end = next((i for i in range(ingredients_start + 1, len(lines)) if not lines[i].strip()), len(lines))
        ingredients = [line.strip() for line in lines[ingredients_start + 1:ingredients_end]]

        # Extract Method
        method_end = next((i for i in range(method_start + 1, len(lines)) if 'Notes:' in lines[i]), len(lines))
        method = [line.strip() for line in lines[method_start + 1:method_end]]

        return recipe_name, ingredients, method
    except Exception as e:
        print(f"Error reading file: {e}")
        return None, None, None

# ========== Main Program ==========
    
def main():
    input_folder = 'parsed'
    output_folder = 'export'

    # Check if the output folders exist, create them if not
    docx_output_folder = os.path.join(output_folder, 'docx')
    pdf_output_folder = os.path.join(output_folder, 'pdf')
    
    for folder in [output_folder, docx_output_folder, pdf_output_folder]:
        if not os.path.exists(folder):
            os.makedirs(folder)

    # Iterate through all .txt files in the 'parsed' folder
    for file_name in os.listdir(input_folder):
        if file_name.endswith(".txt"):
            input_file_path = os.path.join(input_folder, file_name)

            # Read recipe details from the text file
            recipe_name, ingredients, method = read_recipe_from_file(input_file_path)

            # Check if the file was read successfully
            if recipe_name is not None:
                # Create a new Word document with the recipe information
                doc = create_recipe_doc(recipe_name, ingredients, method)

                # Save the document to a file using the recipe name as the file name
                output_doc_path = os.path.join(docx_output_folder, f"{recipe_name.replace(' ', '_')}.docx")
                doc.save(output_doc_path)
                print(f"[OK] EXPORT - Recipe Document '{output_doc_path}' Successfully Created")
            else:
                print(f"[ERROR] EXPORT - Recipe details could not be extracted from {file_name}.")

    # Iterate through all .docx files in the 'export/docx' folder
    for file_name in os.listdir(docx_output_folder):
        if file_name.endswith(".docx"):
            print(f"[OK] EXPORT - Converting Document '{file_name}' to PDF")
            input_file_path = os.path.join(docx_output_folder, file_name)
            output_file_path = os.path.join(pdf_output_folder, file_name.replace(".docx", ".pdf"))
            convert(input_file_path, output_file_path)
            print(f"[OK] EXPORT - PDF Document '{output_file_path}' Successfully Created")

if __name__ == "__main__":
    main()